import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import os
import tempfile
import seaborn as sns
import plotly.express as px
import google.generativeai as genai
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

genai.configure(api_key="AIzaSyDoQSwSMdwX5uCNgFQKwKj63EVBc8VPSFs")

st.set_page_config(
    page_title="Data App",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown(
    """
    <style>
    @media (max-width: 768px) {
        .block-container {
            padding: 1rem;
        }
        .stButton > button {
            width: 100%;
        }
    }
    iframe {
        width: 100% !important;
        height: calc(100vh - 200px) !important;
        border: none;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

if "data" not in st.session_state:
    st.session_state.data = None
if "cleaned_data" not in st.session_state:
    st.session_state.cleaned_data = None
if "visualization_settings" not in st.session_state:
    st.session_state.visualization_settings = {"graph": None}
if "interpretation" not in st.session_state:
    st.session_state.interpretation = None
if "fig" not in st.session_state:
    st.session_state.fig = None
if "docx_buffer" not in st.session_state:
    st.session_state.docx_buffer = None
if "pdf_buffer" not in st.session_state:
    st.session_state.pdf_buffer = None

if "previous_choices" not in st.session_state:
    st.session_state.previous_choices = {
        "graph_type": None,
        "x_axis": None,
        "y_axis": None,
        "category_col": None,
        "numeric_col_for_pie": None,
        "class_column": None,
        "numeric_cols_selected": None
    }

generation_config = {
    "temperature": 1,
    "top_p": 0.95,
    "top_k": 40,
    "max_output_tokens": 8192,
    "response_mime_type": "text/plain",
}

model = genai.GenerativeModel(
    model_name="gemini-1.5-pro",
    generation_config=generation_config,
)

def generate_explanation(prompt):
    try:
        chat_session = model.start_chat(history=[])
        response = chat_session.send_message(prompt)
        return response.text
    except Exception as e:
        return f"An error occurred while generating the explanation: {str(e)}"

def create_docx(interpretation, fig):
    doc = Document()

    title = doc.add_heading("AI Interpretation and Visualization", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    heading = doc.add_heading("Data Insights", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph(interpretation)
    p.style = doc.styles['Normal']
    p_format = p.paragraph_format
    p_format.space_after = Pt(12)
    p_format.line_spacing = 1.5

    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmpfile:
        fig.savefig(tmpfile.name, format="png")
        tmpfile.flush()
        doc.add_picture(tmpfile.name, width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmpfile:
        doc.save(tmpfile.name)
        tmpfile.flush()
        with open(tmpfile.name, "rb") as f:
            docx_buffer = BytesIO(f.read())

    return docx_buffer

def create_pdf(interpretation, fig):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmpfile:
        fig.savefig(tmpfile.name, format="png", dpi=300)
        tmpfile.flush()
        img_path = tmpfile.name

    pdf = FPDF()
    pdf.add_page()

    pdf.set_font("Arial", 'B', 20)
    pdf.cell(0, 5, "AI Interpretation and Visualization", align='C', ln=True)
    pdf.ln(10)

    pdf.set_font("Arial", '', 12)

    lines = interpretation.split('\n')
    for line in lines:
        pdf.multi_cell(0, 5, line)
        pdf.ln(5)

    pdf.add_page()

    pdf.set_font("Arial", 'B', 16)
    pdf.cell(0, 7, "Visualization", ln=True, align='C')
    pdf.ln(10)
    pdf.image(img_path, x=10, y=40, w=180)
    pdf.ln(100)

    pdf_data = pdf.output(dest='S').encode('latin-1', 'replace')
    pdf_buffer = BytesIO(pdf_data)
    return pdf_buffer

def DataCleaning():
    st.title("Data Cleaning")
    st.write("Clean your data here.")

    uploaded_file = st.file_uploader("Choose a file to clean")
    if uploaded_file is not None:
        st.session_state.data = pd.read_csv(uploaded_file)
        st.dataframe(st.session_state.data)

        with st.expander("Summary Statistics"):
            st.write(st.session_state.data.describe())

        with st.expander("Missing Values"):
            st.write(st.session_state.data.isnull().sum())

        with st.expander("Duplicate Rows"):
            st.write(st.session_state.data.duplicated().sum())

        with st.expander("Data Types"):
            st.write(st.session_state.data.dtypes)
    else:
        st.session_state.data = None
        st.session_state.interpretation = None
        st.session_state.fig = None
        st.session_state.docx_buffer = None
        st.session_state.pdf_buffer = None

    if st.button("Clean Data"):
        if st.session_state.data is not None:
            initial_empty_cells = st.session_state.data.isnull().sum().sum()
            st.session_state.cleaned_data = st.session_state.data.dropna(axis=1)
            final_empty_cells = st.session_state.cleaned_data.isnull().sum().sum()
            empty_cells_removed = initial_empty_cells - final_empty_cells
            st.write(
                f"**Data cleaned successfully. {empty_cells_removed} empty cells removed.**"
            )
            st.dataframe(st.session_state.cleaned_data)

            cleaned_csv = st.session_state.cleaned_data.to_csv(index=False).encode('utf-8')
            st.download_button(
                label="📥 Download Cleaned Data as CSV",
                data=cleaned_csv,
                file_name="cleaned_data.csv",
                mime="text/csv",
            )
        else:
            st.error("Please upload a file first.")

def Visualization():
    st.title("Data Visualization with AI Interpretation")
    st.write("Visualize your data and get AI insights here.")

    uploaded_file = st.file_uploader("Choose a file to visualize")
    if uploaded_file is not None:
        st.session_state.data = pd.read_csv(uploaded_file)
        df = st.session_state.data

        graph_type = st.selectbox(
            "Select Graph Type",
            options=["Bar Chart", "Line Chart", "Pie Chart", "Parallel Coordinates", "Area Chart"]
        )
        st.session_state.visualization_settings["graph"] = graph_type

        numeric_columns = df.select_dtypes(include=["number"]).columns.tolist()
        categorical_columns = df.select_dtypes(exclude=["number"]).columns.tolist()

        measure_options = ["Count", "Sum", "Average", "Min", "Max"]

        x_axis = None
        y_axis = None
        category_col = None
        numeric_col_for_pie = None
        class_column = None
        numeric_cols_selected = None
        measure = "Count"

        if graph_type in ["Bar Chart", "Line Chart", "Area Chart"]:
            x_axis = st.selectbox("Select X-axis (Categorical)", options=df.columns)
            y_axis = st.selectbox("Select Y-axis (Numeric)", options=df.columns)

            valid_x = x_axis in df.columns
            valid_y = y_axis in numeric_columns or y_axis == x_axis

            if not valid_x or "Unnamed" in x_axis:
                st.error("Please select a valid X-axis column.")
            if not valid_y or "Unnamed" in y_axis:
                st.error("Please select a valid numeric Y-axis column.")

            if valid_y and y_axis in numeric_columns:
                measure = st.selectbox("Measure Values", options=measure_options)
            else:
                measure = "Count"

        elif graph_type == "Pie Chart":
            category_col = st.selectbox("Select Category Column", options=df.columns)

            numeric_col_for_pie = st.selectbox("Select Numeric Column (optional)", options=["None"]+numeric_columns)
            if numeric_col_for_pie == "None":
                measure = "Count"
                numeric_col_for_pie = None
            else:
                measure = st.selectbox("Measure Values", options=measure_options)

        elif graph_type == "Parallel Coordinates":
            if len(categorical_columns) == 0:
                st.error("You need at least one categorical column for Parallel Coordinates.")
                return
            if len(numeric_columns) < 2:
                st.error("You need at least two numeric columns for Parallel Coordinates.")
                return
            class_column = st.selectbox("Select Categorical Column for Class", options=categorical_columns)
            numeric_cols_selected = st.multiselect("Select Numeric Columns", options=numeric_columns, default=numeric_columns[:2])
            measure = None
        else:
            measure = "Count"

        current_choices = {
            "graph_type": graph_type,
            "x_axis": x_axis,
            "y_axis": y_axis,
            "category_col": category_col,
            "numeric_col_for_pie": numeric_col_for_pie,
            "class_column": class_column,
            "numeric_cols_selected": tuple(numeric_cols_selected) if numeric_cols_selected else None
        }

        if any(st.session_state.previous_choices[key] != current_choices[key] for key in current_choices):
            st.session_state.fig = None
            st.session_state.interpretation = None
            st.session_state.docx_buffer = None
            st.session_state.pdf_buffer = None

        st.session_state.previous_choices = current_choices

        if st.button("Generate Visualization & Interpretation"):
            try:
                df_grouped = None
                fig, ax = plt.subplots(figsize=(10, 10))
                fig.tight_layout()
                fig.subplots_adjust(left=0.1, right=1, top=0.9, bottom=0.1)

                if graph_type in ["Bar Chart", "Line Chart", "Area Chart"]:
                    if measure == "Count":
                        df_grouped = df.groupby(x_axis).size()
                        y_label = "Count"
                    else:
                        if y_axis not in numeric_columns:
                            st.error("Selected Y-axis must be numeric for the chosen measure.")
                            return
                        if measure == "Sum":
                            df_grouped = df.groupby(x_axis)[y_axis].sum()
                        elif measure == "Average":
                            df_grouped = df.groupby(x_axis)[y_axis].mean()
                        elif measure == "Min":
                            df_grouped = df.groupby(x_axis)[y_axis].min()
                        elif measure == "Max":
                            df_grouped = df.groupby(x_axis)[y_axis].max()
                        y_label = f"{measure} of {y_axis}"

                    if graph_type == "Bar Chart":
                        df_grouped.plot(kind="bar", ax=ax)
                        ax.set_title("Bar Chart")
                    elif graph_type == "Line Chart":
                        df_grouped.plot(kind="line", ax=ax)
                        ax.set_title("Line Chart")
                    elif graph_type == "Area Chart":
                        df_grouped.plot(kind="area", ax=ax)
                        ax.set_title("Area Chart")

                    ax.set_xlabel(x_axis)
                    ax.set_ylabel(y_label)

                elif graph_type == "Pie Chart":
                    if category_col is None:
                        st.error("Please select a valid category column for the Pie Chart.")
                        return
                    if measure == "Count":
                        df_grouped = df.groupby(category_col).size()
                    else:
                        if numeric_col_for_pie not in numeric_columns:
                            st.error("Selected numeric column for Pie Chart must be valid and numeric.")
                            return
                        if measure == "Sum":
                            df_grouped = df.groupby(category_col)[numeric_col_for_pie].sum()
                        elif measure == "Average":
                            df_grouped = df.groupby(category_col)[numeric_col_for_pie].mean()
                        elif measure == "Min":
                            df_grouped = df.groupby(category_col)[numeric_col_for_pie].min()
                        elif measure == "Max":
                            df_grouped = df.groupby(category_col)[numeric_col_for_pie].max()

                    df_grouped.plot(kind="pie", ax=ax, autopct='%1.1f%%')
                    ax.set_title("Pie Chart")
                    ax.set_ylabel("")

                elif graph_type == "Parallel Coordinates":
                    if class_column not in categorical_columns:
                        st.error("Class column must be categorical.")
                        return
                    if len(numeric_cols_selected) < 2:
                        st.error("Please select at least two numeric columns.")
                        return

                    subset_df = df[[class_column] + numeric_cols_selected].dropna()
                    pd.plotting.parallel_coordinates(subset_df, class_column, ax=ax)
                    ax.set_title("Parallel Coordinates")

                st.session_state.fig = fig

            except Exception as e:
                st.error(f"An error occurred while creating the plot: {str(e)}")
                return

            if df_grouped is not None:
                grouped_values = df_grouped.reset_index().values.tolist()
            else:
                grouped_values = []

            if graph_type in ["Bar Chart", "Line Chart", "Area Chart"]:
                explanation_prompt = (
                    f"You have created a {graph_type.lower()} using the column '{x_axis}' as the X-axis and "
                    f"applied the measure '{measure}' on the column '{y_axis}'. "
                    f"The aggregated data is as follows (X-axis value, Aggregated result): {grouped_values}. "
                    f"Explain the trends, patterns, and insights visible in this data. "
                    f"Consider how the values vary across different '{x_axis}' categories."
                )

            elif graph_type == "Pie Chart":
                if measure == "Count":
                    explanation_prompt = (
                        f"A pie chart has been created using '{category_col}' as the category dimension, "
                        f"showing the percentage distribution of counts among each category. "
                        f"Here are the values (Category, Count): {grouped_values}. "
                        f"Describe what this data reveals about the relative proportions of the categories."
                    )
                else:
                    explanation_prompt = (
                        f"A pie chart has been created using '{category_col}' as the category dimension. "
                        f"The values represent the {measure.lower()} of '{numeric_col_for_pie}'. "
                        f"Here are the aggregated values (Category, Aggregated result): {grouped_values}. "
                        f"Explain what this data shows about the distribution of these values across the categories."
                    )

            elif graph_type == "Parallel Coordinates":
                sample_values = subset_df.head(100).to_dict(orient='records')
                explanation_prompt = (
                    f"A parallel coordinates plot has been created using '{class_column}' as the class column "
                    f"and {numeric_cols_selected} as the numeric features. "
                    f"Here is a sample of the data used:\n{sample_values}\n"
                    f"Interpret how different classes differ across these numeric dimensions and describe any patterns you observe."
                )

            with st.spinner("Generating AI interpretation, please wait..."):
                explanation = generate_explanation(explanation_prompt)

            st.session_state.interpretation = explanation

    else:
        st.session_state.data = None
        st.session_state.interpretation = None
        st.session_state.fig = None
        st.session_state.docx_buffer = None
        st.session_state.pdf_buffer = None

    if st.session_state.fig is not None and st.session_state.interpretation is not None:
        col1, col2 = st.columns([2, 1])
        with col1:
            st.pyplot(st.session_state.fig)
        with col2:
            st.subheader("AI Interpretation")
            st.write(st.session_state.interpretation)

        format_choice = st.radio("Select Document Format:", ("DOCX", "PDF"), index=0)

        if format_choice == "DOCX" and st.session_state.pdf_buffer is not None:
            st.session_state.pdf_buffer = None
        elif format_choice == "PDF" and st.session_state.docx_buffer is not None:
            st.session_state.docx_buffer = None

        if st.button("Generate Document"):
            with st.spinner("Generating your document..."):
                if format_choice == "DOCX":
                    st.session_state.docx_buffer = create_docx(st.session_state.interpretation, st.session_state.fig)
                    st.session_state.pdf_buffer = None
                else:
                    st.session_state.pdf_buffer = create_pdf(st.session_state.interpretation, st.session_state.fig)
                    st.session_state.docx_buffer = None
            st.success("Document generated successfully!")

        if st.session_state.docx_buffer is not None and format_choice == "DOCX":
            st.download_button(
                label="📥 Download DOCX",
                data=st.session_state.docx_buffer,
                file_name="interpretation_and_visualization.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        elif st.session_state.pdf_buffer is not None and format_choice == "PDF":
            st.download_button(
                label="📥 Download PDF",
                data=st.session_state.pdf_buffer,
                file_name="interpretation_and_visualization.pdf",
                mime="application/pdf",
            )

def main():
    st.sidebar.title("Navigation")
    page = st.sidebar.selectbox("Go to", ["Data Cleaning", "Data Visualization"])

    if st.session_state.data is None:
        st.session_state.interpretation = None
        st.session_state.fig = None
        st.session_state.docx_buffer = None
        st.session_state.pdf_buffer = None

    if page == "Data Cleaning":
        DataCleaning()
    elif page == "Data Visualization":
        Visualization()

if __name__ == "__main__":
    main()
